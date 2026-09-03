from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"D:\StratonAlly\Code\spareautomation")
OUTPUT = ROOT / "client-documents" / "Spares Automation - Complete Client Operations Manual.docx"

NAVY = "172235"
ORANGE = "E6602E"
LIGHT_ORANGE = "FFF1EB"
LIGHT_BLUE = "EAF0F6"
LIGHT_GREY = "F2F4F6"
MID_GREY = "667085"
WHITE = "FFFFFF"
GREEN = "DFF3E4"
RED = "FCE8E6"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, value: str, *, bold=False, color=None, size=9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(value))
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def table(doc: Document, headers: list[str], rows: list[list[str]], widths=None, font_size=8.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    t.autofit = True
    for i, header in enumerate(headers):
        shade(t.rows[0].cells[i], NAVY)
        set_cell_text(t.rows[0].cells[i], header, bold=True, color=WHITE, size=8.5)
    for row_index, values in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(values):
            if row_index % 2:
                shade(cells[i], LIGHT_GREY)
            set_cell_text(cells[i], value, size=font_size)
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def callout(doc: Document, title: str, text: str, kind="info") -> None:
    fills = {"info": LIGHT_BLUE, "warning": LIGHT_ORANGE, "success": GREEN, "danger": RED}
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    shade(cell, fills[kind])
    set_cell_margins(cell, top=160, start=180, bottom=160, end=180)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title.upper())
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(ORANGE if kind in {"warning", "danger"} else NAVY)
    p2 = cell.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(0)
    for run in p2.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def bullets(doc: Document, items: list[str], level=0) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.add_run(item)


def steps(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def heading(doc: Document, text: str, level=1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        p.add_run(bold_lead).bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def page_break(doc: Document) -> None:
    doc.add_page_break()


doc = Document()
doc.core_properties.title = "Spares Automation Complete Client Operations Manual"
doc.core_properties.subject = "Website, Shopify catalogue, CMS, orders, quotes, deployment and maintenance"
doc.core_properties.author = "Spares Automation project team"
doc.core_properties.keywords = "Spares Automation, Shopify, CMS, products, operations"
doc.core_properties.comments = "Prepared from the current application source and operating configuration; secrets excluded."

section = doc.sections[0]
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)
section.header_distance = Cm(0.8)
section.footer_distance = Cm(0.8)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal.font.size = Pt(9.5)
normal.font.color.rgb = RGBColor.from_string(NAVY)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.08

for name, size, color, before, after in (
    ("Title", 34, NAVY, 0, 8),
    ("Heading 1", 22, NAVY, 15, 6),
    ("Heading 2", 15, ORANGE, 11, 4),
    ("Heading 3", 11.5, NAVY, 8, 3),
    ("Heading 4", 10, MID_GREY, 6, 2),
):
    st = styles[name]
    st.font.name = "Aptos Display"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for name in ("List Bullet", "List Bullet 2", "List Number"):
    styles[name].font.name = "Aptos"
    styles[name].font.size = Pt(9.5)
    styles[name].paragraph_format.space_after = Pt(2.5)

if "Caption Small" not in styles:
    cap = styles.add_style("Caption Small", WD_STYLE_TYPE.PARAGRAPH)
    cap.font.name = "Aptos"
    cap.font.size = Pt(8)
    cap.font.italic = True
    cap.font.color.rgb = RGBColor.from_string(MID_GREY)

# Header/footer
header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = header.add_run("SPARES AUTOMATION  /  CLIENT OPERATIONS MANUAL")
run.font.name = "Aptos"
run.font.size = Pt(7.5)
run.font.bold = True
run.font.color.rgb = RGBColor.from_string(MID_GREY)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run("Confidential client operating document  •  ")
r.font.name = "Aptos"
r.font.size = Pt(7.5)
r.font.color.rgb = RGBColor.from_string(MID_GREY)
add_field(footer, "PAGE", "1")

# Make Word refresh fields (including TOC) on open.
settings = doc.settings._element
update_fields = OxmlElement("w:updateFields")
update_fields.set(qn("w:val"), "true")
settings.append(update_fields)

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(40)
r = p.add_run("SPARES\nAUTOMATION")
r.font.name = "Aptos Display"
r.font.size = Pt(17)
r.font.bold = True
r.font.color.rgb = RGBColor.from_string(ORANGE)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(55)
p.paragraph_format.space_after = Pt(12)
r = p.add_run("Complete Client\nOperations Manual")
r.font.name = "Aptos Display"
r.font.size = Pt(36)
r.font.bold = True
r.font.color.rgb = RGBColor.from_string(NAVY)

p = doc.add_paragraph()
r = p.add_run("Website • Shopify catalogue • CMS • Orders • Quotes • Maintenance")
r.font.name = "Aptos"
r.font.size = Pt(13)
r.font.color.rgb = RGBColor.from_string(MID_GREY)

doc.add_paragraph("\n")
table(doc, ["Document", "Details"], [
    ["Version", "1.0"],
    ["Prepared", "1 September 2026"],
    ["System", "Spares Automation website and connected services"],
    ["Website", "https://spares-automation.co.uk"],
    ["Audience", "Business owners, catalogue managers, sales staff, returns staff and technical operators"],
    ["Source reviewed", "Current repository implementation at commit ea03d3e (secrets excluded)"],
], widths=[4, 12], font_size=9)

callout(doc, "Purpose", "This is the operating source of truth for the system currently implemented. Shopify and hosting interfaces can change their labels over time; when a screen differs, follow the ownership rules and field names in this manual.", "info")

page_break(doc)
heading(doc, "How to use this manual", 1)
para(doc, "Start with Sections 1–3 to understand what each system owns. Product and catalogue managers should use Sections 5–7. Sales and support teams should use Sections 8–11. Technical operators should use Sections 14–18. Checklists and reference tables are in the appendices.")
callout(doc, "Security", "No passwords, API keys, database connection strings or other secret values are included. Obtain credentials through the organisation’s approved password manager or system owner. Never send the .env file to a client or vendor.", "warning")

heading(doc, "Document control", 2)
table(doc, ["When", "Owner", "Required update"], [
    ["After a website release", "Developer / technical owner", "Update changed routes, workflows, configuration and known limitations."],
    ["After a Shopify catalogue policy change", "Catalogue owner", "Update product-field rules, collection assignments and publishing steps."],
    ["After a CMS process change", "Operations owner", "Update status meanings, response targets and staff responsibilities."],
    ["At least quarterly", "Business owner", "Review contacts, legal copy, access list, backups and recovery checks."],
], widths=[3, 4, 10])

heading(doc, "Table of contents", 1)
toc = doc.add_paragraph()
add_field(toc, 'TOC \\o "1-3" \\h \\z \\u', "Right-click and choose Update Field if the table does not populate automatically.")

page_break(doc)
heading(doc, "1. System overview", 1)
para(doc, "Spares Automation is an industrial-parts ecommerce and enquiry website built with TanStack Start and React. It presents a Shopify catalogue, supports direct Shopify checkout, creates sales quotations as Shopify draft orders, provides Shopify-backed customer accounts, and records non-order enquiries in a lightweight internal CMS backed by PostgreSQL.")

heading(doc, "1.1 The three operating systems", 2)
table(doc, ["System", "Owns", "Does not own"], [
    ["Shopify Admin", "Products, variants, SKUs, prices, inventory, images, manual collections, publication to Headless, checkout, orders, customers, draft orders, customer marketing consent.", "Website page copy, CMS notes, CMS return workflow, homepage images."],
    ["Website CMS at /admin", "Part enquiries, credit applications, return requests, order-tracking requests, resource requests, product questions, unsubscribe requests, attachments, internal notes and statuses.", "Products, carts, checkout, orders, invoices, general page content."],
    ["Application code / deployment", "Navigation, homepage and category presentation, category contract, page copy, legal pages, contact details, styling, forms, integrations, environment settings and database migrations.", "Day-to-day product or order management."],
], widths=[3, 7, 7], font_size=8.3)

callout(doc, "Most important rule", "Orders always stay in Shopify. The CMS never creates or manages orders. Approved credit applications can create or update a tagged Shopify customer, after which staff raise draft orders or invoices in Shopify.", "warning")

heading(doc, "1.2 Information flow", 2)
table(doc, ["Event", "Where it starts", "Where it ends"], [
    ["Product published", "Shopify product + Headless channel + collections", "Automatically appears in the website catalogue after Shopify caching refreshes."],
    ["Online checkout", "Website cart", "Shopify checkout and Shopify order."],
    ["Quote submitted", "Website quote list", "Shopify draft order tagged “Website quote request” and “Awaiting sales review”."],
    ["Public form submitted", "Website form", "PostgreSQL CMS submission; sales-desk email is attempted."],
    ["Return status changed", "CMS submission detail", "CMS status is updated and a customer email is attempted."],
    ["Credit application synced", "CMS credit application", "Shopify customer is created/updated and tagged as a credit account; CMS is marked approved."],
], widths=[4, 5, 8])

heading(doc, "2. Roles and responsibilities", 1)
table(doc, ["Role", "Primary responsibilities", "Access needed"], [
    ["Business owner", "Approves catalogue policy, legal wording, customer service rules, access and retention policy.", "Shopify owner/admin; hosting oversight."],
    ["Catalogue manager", "Creates and edits products, media, variants, inventory, collections, resources and publication.", "Shopify Products and Collections."],
    ["Sales desk", "Reviews quotes, turns draft orders into invoices/orders, answers enquiries and handles credit customers.", "Shopify Orders/Customers/Draft orders; CMS."],
    ["Returns/support staff", "Reviews return and support requests, downloads files, adds internal notes and updates statuses.", "CMS; Shopify order lookup."],
    ["CMS administrator", "Controls initial admin credentials and coordinates staff access changes.", "CMS; deployment environment or database support."],
    ["Technical operator", "Deploys releases, manages environment values, database/upload backups, email and Shopify integration health.", "Coolify/host, repository, Shopify app/API configuration."],
], widths=[3, 9, 5], font_size=8.3)

heading(doc, "2.1 Access checklist", 2)
bullets(doc, [
    "Use named Shopify staff accounts with the least permissions required; do not share the store-owner login.",
    "Use the organisation password manager for Shopify, CMS, Coolify, email provider and database credentials.",
    "Remove access promptly when a staff member leaves or changes role.",
    "The CMS admin session lasts up to eight hours. Sign out on shared devices.",
    "The current CMS has no staff-management or password-reset screen. Adding staff, changing a CMS password or revoking a CMS user requires technical/database assistance.",
])

heading(doc, "3. Daily operating map", 1)
table(doc, ["I need to…", "Use", "Action"], [
    ["Add or edit a product", "Shopify Admin", "Products → create/edit → publish to Headless → assign collections."],
    ["Change price, SKU, stock or images", "Shopify Admin", "Edit the product or use Shopify bulk editor."],
    ["Make a product appear in a category", "Shopify Admin", "Assign the manual main-category and product-line collections."],
    ["Process an online order", "Shopify Admin", "Orders; follow Shopify fulfilment, payment and refund workflows."],
    ["Review a quote", "Shopify Admin", "Draft orders; find the Website quote request tag."],
    ["Review a customer enquiry or return", "Website CMS", "Open /admin, filter the queue, open the record, add notes and set status."],
    ["Approve a credit application", "CMS then Shopify", "Verify data, sync once to Shopify, then manage orders/invoices in Shopify."],
    ["Change homepage text, legal text or navigation", "Developer release", "Request a code change; this is not editable in Shopify or the CMS."],
    ["Change category names/handles", "Developer + Shopify", "Coordinate a catalogue-contract code change and matching collections."],
], widths=[5, 3, 9], font_size=8.3)

heading(doc, "4. Customer-facing website", 1)
heading(doc, "4.1 Main customer journeys", 2)
bullets(doc, [
    "Browse the home page, six catalogue families and their product lines.",
    "Search globally by query or filter the catalogue by collection, availability and sort order.",
    "Open product details, select a variant and quantity, then add to cart or add to quote.",
    "Complete secure checkout in Shopify, or submit a quote for staff review.",
    "Register or sign in to a Shopify customer account, view linked order history when available, and follow Shopify order-status links.",
    "Submit part-identification enquiries, credit applications, returns, tracking requests, resource requests, product questions and unsubscribe requests.",
    "Browse the resource library assembled from product metafields.",
])

heading(doc, "4.2 Catalogue filters", 2)
table(doc, ["Control", "Options / behavior"], [
    ["Category", "All products, any of the six main collections, or a product-line collection."],
    ["Availability", "All Stock or Available. This is a storefront filter; correct Shopify inventory/availability is essential."],
    ["Sort", "Newest, A–Z, Price Low, Price High."],
    ["Pagination", "Loads 48 products initially; “Load more products” uses Shopify cursor pagination."],
    ["URL persistence", "Category, availability and sort are stored in the URL so a filtered view can be shared/bookmarked."],
], widths=[4, 13])

heading(doc, "4.3 Cart versus quote", 2)
table(doc, ["Path", "Result", "Staff action"], [
    ["Cart → Checkout", "Customer continues to secure Shopify checkout. Delivery is calculated there; prices are shown excluding VAT on the website cart.", "Manage the resulting Shopify order."],
    ["Build a Quote", "Browser-local quote list is submitted to Shopify as a draft order. Final price, VAT, availability, delivery and terms are not promised by the website.", "Review the draft order, adjust it and contact the customer."],
], widths=[4, 8, 5])

heading(doc, "5. Shopify catalogue management", 1)
callout(doc, "Source of truth", "Products and catalogue membership are managed in Shopify. A published product assigned to the correct manual collections appears on the website automatically; no website code edit or product tag is required for ordinary product assignment.", "success")

heading(doc, "5.1 Add a product", 2)
steps(doc, [
    "In Shopify Admin, open Products and choose Add product.",
    "Enter a clear product title and customer-facing description. Use the rich description for specifications and application notes.",
    "Add good-quality product images. Set meaningful alt text for accessibility and search engines.",
    "Set Vendor (used as the fallback brand) and Product type (also used to group resources).",
    "Create the required variants. For every variant, enter a unique SKU/part number, price, optional compare-at price, inventory and availability.",
    "Set the product status to Active.",
    "In Publishing / Sales channels and apps, make the product available to the Headless sales channel used by this website.",
    "In Collections, assign one main storefront category and at least one specific product-line collection where appropriate.",
    "Add optional technical/resource metafields described in Section 6.",
    "Save, open the website category and product page, and test the correct variant, availability, cart and quote actions.",
])

heading(doc, "5.2 Required product quality standard", 2)
table(doc, ["Field", "Required practice", "Website use"], [
    ["Title", "Concise customer-facing name; distinguish similar models.", "Cards, search results, cart, quote and product heading."],
    ["Handle", "Keep stable once public. If changed, old links may break unless redirected.", "Product URL /products/{handle}."],
    ["Description", "Explain purpose, compatibility and limitations; do not make unapproved delivery claims.", "Product detail and search context."],
    ["Images", "Use clean, sufficiently large images; set accurate alt text; avoid text-heavy artwork.", "Cards, gallery, cart and quote."],
    ["Vendor", "Use the actual manufacturer/brand consistently.", "Fallback Brand when custom.brand is absent."],
    ["Product type", "Use a consistent family/category label.", "Resource-library grouping fallback."],
    ["Variants", "Use variants only for genuine purchasable options; make titles understandable.", "Variant selector and cart/quote lines."],
    ["SKU", "Unique part number on every variant.", "Product details, cart, quote and identification."],
    ["Price", "Correct currency and ex-VAT trade price policy.", "Cards, details, cart and indicative quote subtotal."],
    ["Inventory", "Maintain quantity/availability or the chosen Shopify inventory policy.", "Available filter and add-to-cart state."],
    ["Status/channel", "Active and available to Headless.", "Both are required for storefront visibility."],
    ["Collections", "Main category plus product line.", "Navigation and catalogue filtering."],
], widths=[3, 7, 7], font_size=8.1)

heading(doc, "5.3 Edit or bulk-update products", 2)
bullets(doc, [
    "For one product: Products → open product → edit → Save.",
    "For several products: select them in Products and use Bulk edit or Add to collections.",
    "After a bulk change, inspect representative products on the website. Shopify storefront caching may take a short time to refresh.",
    "Do not rename required collection handles casually. Handles form an integration contract with the website.",
])

heading(doc, "5.4 Product visibility troubleshooting", 2)
steps(doc, [
    "Confirm the product status is Active.",
    "Confirm it is available on the Headless sales channel/publication.",
    "Confirm it belongs to the correct main manual collection and product-line manual collection.",
    "Confirm the variant is available for sale and inventory policy is correct.",
    "Confirm the product/collection handle matches the expected value and wait briefly for Shopify caching.",
    "Check the All Products view. If absent there too, escalate to the Shopify/API technical owner.",
])

heading(doc, "6. Product resources and technical metafields", 1)
para(doc, "The product page and /resources library read metafields in the custom namespace. Only HTTPS links are accepted by the website. Unsafe, malformed or non-HTTPS values are ignored.")
table(doc, ["Metafield", "Expected value", "Where shown"], [
    ["custom.brand", "Plain text. Overrides Vendor as the displayed brand.", "Product technical details."],
    ["custom.mpn_range", "Plain text manufacturer part number/range. Falls back to first variant SKU.", "Product technical details."],
    ["custom.setup_video_url", "One HTTPS URL.", "Product video area and resource library."],
    ["custom.video_guide", "JSON object: {\"text\":\"Setup guide\",\"url\":\"https://…\"}", "Named video link and resource library."],
    ["custom.pdf_guide", "JSON object: {\"text\":\"Product guide\",\"url\":\"https://…\"}", "Named PDF/document link and resource library."],
    ["custom.datasheets", "HTTPS URL, JSON object, or JSON array of {label,url} objects.", "Datasheet links and resource library."],
    ["custom.manuals", "HTTPS URL, JSON object, or JSON array of {label,url} objects.", "Manual links and resource library."],
], widths=[4, 8, 5], font_size=8.1)

heading(doc, "6.1 Resource-link example", 2)
para(doc, '[{"label":"Installation manual","url":"https://files.example.com/manual.pdf"},{"label":"Technical datasheet","url":"https://files.example.com/datasheet.pdf"}]')
bullets(doc, [
    "Host files at stable HTTPS URLs accessible without an internal login.",
    "Use specific labels rather than “Resource 1”.",
    "Test every link in a private browser window.",
    "Videos require customer consent before third-party YouTube content loads; this supports privacy/cookie behavior.",
    "Products without resource metafields are omitted from the resource library.",
])

heading(doc, "7. Catalogue collections", 1)
para(doc, "All storefront categories are Shopify manual collections. The collection handle must exactly match the website contract below. A product normally belongs to one main collection and one or more product-line collections.")

categories = [
    ("Asphalt / Blacktop Spares", "asphalt", "Aggregate Feeding", "asphalt-aggregate-feeding"),
    ("", "", "Burner / Drying", "burner-drying"),
    ("", "", "Bitumen", "bitumen"),
    ("", "", "Hot Storage / Silos", "hot-storage-and-silos"),
    ("", "", "Baghouse", "baghouse"),
    ("", "", "Mixing Tower", "mixing-tower"),
    ("Concrete Spares", "concrete", "Aggregate Feeding", "concrete-aggregate-feeding"),
    ("", "", "Cement / Material Silos", "cement-material-silos"),
    ("", "", "Additive System", "additive-system"),
    ("", "", "Water Controls", "water-controls"),
    ("", "", "Air Controls", "air-controls"),
    ("", "", "Automation Controls", "automation-controls"),
    ("Packing Machinery", "packing", "Automation / Sensors", "automation-sensors"),
    ("", "", "Bag Placement", "bag-placement"),
    ("", "", "Filling", "filling"),
    ("", "", "Discharge & Palletising", "discharge-palletising"),
    ("Automation & Drives", "automation", "Contactors", "contactors"),
    ("", "", "Sensors", "sensors"),
    ("", "", "Buttons / Switches", "buttons-switches"),
    ("", "", "Inverter Drives", "inverter-drives"),
    ("Home Automation and Controls", "home-controls", "Lighting", "lighting"),
    ("", "", "Security", "security"),
    ("Control Panels & Software", "control-panels-software", "Control Panels", "control-panels"),
    ("", "", "Software Design and Programming", "software-design-programming"),
]
table(doc, ["Main collection", "Main handle", "Product line", "Product-line handle"], [list(x) for x in categories], widths=[4.5, 4, 4.5, 4], font_size=7.6)

heading(doc, "7.1 One-time collection synchronisation (technical owner)", 2)
para(doc, "The repository contains a safe preview/apply tool that creates missing collections, aligns titles and publishes them to the Headless publication. It does not assign products to collections.")
steps(doc, [
    "Install and authenticate Shopify CLI to the permanent store domain with read_products, write_products, read_publications and write_publications scopes, unless a suitable Admin API token is configured.",
    "Run npm run shopify:sync-collections for a dry-run preview.",
    "Review CREATE, PUBLISH and RENAME output. If multiple Headless publications exist, configure the requested publication ID.",
    "Run npm run shopify:sync-collections -- --apply only after approving the preview.",
    "Assign products manually or in bulk in Shopify; the tool does not infer product membership.",
])
callout(doc, "Change control", "Renaming a category or changing a handle is a coordinated developer + Shopify change. Do not create a second near-duplicate collection to work around a mismatch.", "warning")

heading(doc, "8. Orders, checkout and fulfilment", 1)
para(doc, "The website creates a Shopify cart and sends the customer to Shopify checkout. Payment, tax, shipping, order status, fulfilment, cancellations and refunds are Shopify responsibilities.")
heading(doc, "8.1 Order operating procedure", 2)
steps(doc, [
    "Monitor Shopify Orders and the sales mailbox according to the business response schedule.",
    "Open the order; verify payment state, customer details, delivery address, tax and line items.",
    "Confirm stock and fulfil according to the organisation’s warehouse process.",
    "Record tracking and mark fulfilment in Shopify so Shopify’s customer notifications and status page remain accurate.",
    "Handle cancellations/refunds in Shopify, following approved returns and finance policy.",
])
callout(doc, "Configuration validation", "Before public launch, test Shopify Markets, shipping, tax, checkout, payment, customer email and fulfilment with real staging orders. The repository cannot prove the live Shopify business configuration.", "warning")

heading(doc, "9. Quotes and credit accounts", 1)
heading(doc, "9.1 Quote flow", 2)
bullets(doc, [
    "A customer can add a product variant to a browser-local quote list, or copy the current cart into that list.",
    "Quantities are limited to 1–99 per quote item and a submitted quote contains 1–50 lines.",
    "Guest customers enter first name, last name and email; company, phone and additional information are optional.",
    "Signed-in customer details are prefilled when Shopify provides them.",
    "Submission creates a Shopify draft order, not a completed order. It is tagged Website quote request and Awaiting sales review.",
    "The website clears the local quote after success and displays the Shopify draft-order reference.",
])

heading(doc, "9.2 Sales processing", 2)
steps(doc, [
    "In Shopify Admin, open Draft orders and find the Website quote request / Awaiting sales review tags or the reference supplied by the customer.",
    "Verify customer, products, variants, quantities, indicative prices and notes.",
    "Confirm price, VAT, availability, delivery, payment terms and any substitutions with the customer.",
    "For an approved credit customer, apply the agreed account terms; otherwise follow the standard payment process.",
    "Send the invoice/checkout request or complete the draft order using the approved Shopify workflow.",
    "Remove or replace the Awaiting sales review tag according to the team’s internal convention.",
])

heading(doc, "9.3 Credit-application handoff", 2)
steps(doc, [
    "Open the credit application in /admin and independently verify all business, trade-reference and bank information.",
    "Add an internal note recording the approval basis. Do not use a status change as a substitute for credit checks.",
    "Choose Sync to Shopify once. This creates or updates the customer by email, preserves existing tags and adds credit-account and approved-application tags plus reference/metafield data.",
    "The CMS records the Shopify customer ID and sync date and marks the application Approved.",
    "Manage credit limits, draft orders, invoices and payments in Shopify or the organisation’s finance system.",
])
callout(doc, "One-way action", "A credit application can be synced only once from the CMS. The CMS has no undo button. Correct an incorrect customer/tag/metafield directly in Shopify and record the correction in an internal CMS note.", "danger")

heading(doc, "10. CMS access and dashboard", 1)
heading(doc, "10.1 Sign in", 2)
steps(doc, [
    "Open https://spares-automation.co.uk/admin. Unauthenticated users are redirected to /admin/login.",
    "Enter the CMS staff email and password supplied securely by the system owner.",
    "After sign-in, the Submissions dashboard opens. The session expires after up to eight hours.",
    "Use Sign out when finished, particularly on a shared device.",
])
para(doc, "The first administrator is seeded automatically on first boot only when the staff table is empty and the deployment includes ADMIN_SEED_EMAIL and ADMIN_SEED_PASSWORD. Changing those values later does not replace an existing account.")

heading(doc, "10.2 Dashboard controls", 2)
table(doc, ["Control", "Behavior"], [
    ["Status counters", "Totals for New, In review, Approved, Rejected and Completed across active submission types."],
    ["Search", "Matches reference, contact email, contact name or company (case-insensitive)."],
    ["Type filter", "Part inquiry, Credit account, Return request, Order tracking, Resource request, Product question or Unsubscribe."],
    ["Status filter", "Any of the five statuses."],
    ["Pagination", "Newest first, 25 submissions per page."],
    ["Sync column", "Shows whether a credit application has been synced to Shopify."],
], widths=[4, 13])

heading(doc, "10.3 Recommended queue routine", 2)
steps(doc, [
    "At the start of each shift, review the New counter and filter to New.",
    "Open the oldest business-critical items first; verify contact details and captured fields.",
    "Use Mark reviewed to record ownership, or set In review when active work begins.",
    "Add an internal note for every meaningful call, decision, Shopify reference or promised follow-up.",
    "Set the final status only when the business outcome is clear. Recheck New and In review before ending the shift.",
])

heading(doc, "11. CMS submissions in detail", 1)
table(doc, ["Type", "Origin / captured information", "Typical owner"], [
    ["Part inquiry", "Part number, description, name, email, phone; optional JPEG/PNG/WebP photo up to 8 MB.", "Sales / product identification."],
    ["Credit account", "Company, address, type, registration, contacts, two trade references, bank details and signature fields.", "Credit control / sales."],
    ["Return request", "Order, contact, up to 20 items, quantities, reasons, desired outcome, collection address and notes.", "Returns team."],
    ["Order tracking", "Reference/order details, email and message.", "Customer service / fulfilment."],
    ["Resource request", "Product/reference, email and requested document details.", "Technical sales."],
    ["Product question", "Product/reference, email, question and optional supported attachment up to 10 MB.", "Technical sales."],
    ["Unsubscribe", "Reference/email and details; also attempts to update Shopify email-marketing consent.", "Marketing/customer service."],
], widths=[3, 10, 4], font_size=8.2)

heading(doc, "11.1 Submission detail screen", 2)
bullets(doc, [
    "Submission details: contact, company, reference, received time and reviewer.",
    "Captured fields: type-specific form values.",
    "Attachments: images open in a new window; other supported documents download. Files are protected behind staff authentication and are not public URLs.",
    "Notes: staff-only notes, newest first, showing author and time; maximum 4,000 characters each.",
    "Status panel: New, In review, Approved, Rejected or Completed, plus Mark reviewed.",
    "Shopify sync panel: shown only for credit-account applications.",
])

heading(doc, "11.2 Status meanings", 2)
table(doc, ["Status", "Use it when", "Customer communication"], [
    ["New", "Received but not yet owned/reviewed.", "Return requests receive an initial acknowledgement at submission time if email is configured."],
    ["In review", "A staff member is checking or waiting for information.", "Changing a return to this status attempts a return-update email."],
    ["Approved", "The requested action has been accepted.", "Changing a return attempts an approval email. Credit sync also marks the application Approved."],
    ["Rejected", "The request cannot be accepted in its current form.", "Changing a return attempts a rejection email; staff should add a clear internal note."],
    ["Completed", "All required action and communication are finished.", "Changing a return attempts a completion email."],
], widths=[3, 7, 7], font_size=8.2)
callout(doc, "Return emails", "Every actual status change on a return request attempts to email the customer. Check the selected status before clicking. Email failure does not roll back the database change; monitor provider/logs and contact the customer manually when necessary.", "warning")

heading(doc, "11.3 Attachment rules", 2)
table(doc, ["Form", "Allowed", "Limit"], [
    ["Part inquiry", "JPEG, PNG, WebP", "8 MB"],
    ["Product question", "JPEG, PNG, WebP, GIF, PDF, TXT, CSV, DOC/DOCX, XLS/XLSX, PPT/PPTX", "10 MB"],
], widths=[4, 9, 4])
para(doc, "The server validates both the declared MIME type and file signature, sanitises original filenames and stores files in private persistent storage. If the database record exists but the file volume is lost, the CMS will show the attachment but cannot read it.")

heading(doc, "11.4 Email behavior", 2)
bullets(doc, [
    "Resend is preferred when RESEND_API_KEY is configured; SMTP is the fallback.",
    "A form submission is saved even when email is unavailable or rejected. Email failure is logged and does not make the form fail.",
    "Sales-desk notifications go to SALES_DESK_EMAIL; sender identity comes from MAIL_FROM.",
    "Return submission acknowledgements and later return-status emails go to the customer.",
    "Unsubscribe also makes a best-effort Shopify marketing-consent update; failure does not discard the CMS request.",
])

heading(doc, "11.5 CMS limitations", 2)
callout(doc, "Current limitations", "There is no CMS screen to add/remove staff, reset passwords, edit page content, export submissions, delete submissions, define automated assignments, enforce status transitions or resend failed email. Treat these as controlled technical/operational processes, not hidden buttons.", "info")

heading(doc, "12. Customer accounts", 1)
bullets(doc, [
    "Registration and login use Shopify’s Storefront customer flow; the website stores the customer access token only in an encrypted, HTTP-only session cookie.",
    "Customer sessions last up to 14 days, subject to Shopify token expiry.",
    "Password reset requests are sent through Shopify and deliberately use a neutral message so account existence is not disclosed.",
    "The account page displays identity details, recent Shopify orders when available, totals/statuses, return links and Shopify order-status links.",
    "Customers tagged credit account or credit-account receive credit-account quote guidance.",
    "If order history is unavailable through the account API, customers are directed to the tracking-request page.",
])

heading(doc, "13. Website content management", 1)
callout(doc, "No general content CMS", "The /admin area is a submissions workflow only. It cannot edit homepage sections, navigation, category names, contact details, images, legal/policy pages, SEO metadata, robots.txt or sitemap.xml.", "warning")
table(doc, ["Content", "Current owner / location", "How to change"], [
    ["Products, product descriptions and media", "Shopify", "Edit and publish in Shopify."],
    ["Collections and collection descriptions", "Shopify, constrained by code-defined handles", "Edit descriptions in Shopify; coordinate name/handle changes with developer."],
    ["Homepage/category artwork", "Application assets", "Developer replaces asset and deploys. Preferred landscape 4:3 or 16:10, at least 1200 px wide."],
    ["Navigation/category structure", "Application code", "Developer change and deployment."],
    ["Phone, email, WhatsApp, location, hours and canonical domain", "Application configuration in src/lib/site.ts", "Developer change; also review robots.txt and sitemap.xml for domain changes."],
    ["Terms, privacy, returns, delivery, cookies, disclaimer", "Application route content", "Supply approved copy to developer, then deploy."],
    ["Product PDFs/manuals/videos", "Shopify product metafields", "Catalogue manager updates HTTPS resource links."],
], widths=[4.5, 5.5, 7], font_size=8.1)

heading(doc, "14. Deployment and hosting", 1)
para(doc, "The production design uses Docker Compose on Coolify: a private PostgreSQL 16 service and a Node 22 application service exposed internally on port 80. Coolify’s proxy owns the public domain and TLS.")
table(doc, ["Component", "Persistence / behavior"], [
    ["Application container", "Built from the repository; runs compiled Nitro/TanStack output as a non-root node user."],
    ["PostgreSQL", "cms_db named volume at /var/lib/postgresql/data; private, no public domain."],
    ["Uploads", "cms_data named volume mounted at /app/data; UPLOAD_DIR is /app/data/uploads."],
    ["Migrations", "Drizzle SQL migrations run automatically and idempotently at boot, retrying up to five times."],
    ["Initial admin", "Seeded after successful migration only when no staff exists."],
    ["Liveness", "/health returns 200 independently of downstream services."],
    ["Readiness", "/ready checks PostgreSQL and returns 503 when unavailable."],
], widths=[4, 13])

heading(doc, "14.1 Production deployment procedure", 2)
steps(doc, [
    "Review and merge an approved release; ensure lint, typecheck, browser tests and build pass.",
    "In Coolify, use the repository Docker Compose resource and route the public domain to the app service on container port 80. Do not expose or assign a domain to db.",
    "Set required runtime environment variables securely. Never commit production values.",
    "Deploy and monitor build/startup logs for migration, admin seed, email and Shopify configuration errors.",
    "Check /health and /ready, then smoke-test home, catalogue, a product, cart, quote, customer login and CMS login.",
    "Submit a controlled test enquiry; verify CMS storage and email. Where safe, test a quote and remove the test draft order in Shopify.",
])

heading(doc, "14.2 Local Docker operation", 2)
bullets(doc, [
    "docker compose up --build starts the bundled PostgreSQL and app services.",
    "The local override publishes the app at http://localhost:${APP_PORT:-8080}.",
    "Named volumes survive rebuilds. Removing volumes deletes CMS/database or upload data and must be treated as destructive.",
])

heading(doc, "15. Configuration reference", 1)
config_rows = [
    ["DATABASE_URL", "Required for CMS/forms/admin", "PostgreSQL connection used by the running app."],
    ["APP_DATABASE_URL", "Optional", "Use an external managed PostgreSQL URL instead of bundled db in Compose."],
    ["POSTGRES_USER / POSTGRES_DB / SERVICE_PASSWORD_POSTGRES", "Compose database", "Bundled PostgreSQL settings; password must be secret."],
    ["ADMIN_SEED_EMAIL / ADMIN_SEED_PASSWORD", "First CMS admin", "Used only when no staff user exists."],
    ["APP_SESSION_SECRET", "Required", "Encrypts/signs admin and customer session cookies; use a random value of at least 32 bytes."],
    ["RESEND_API_KEY", "Recommended", "Preferred transactional email provider."],
    ["SMTP_HOST / PORT / USER / PASS", "Email fallback", "Used when Resend is absent; port 465 enables secure SMTP."],
    ["MAIL_FROM / SALES_DESK_EMAIL", "Email routing", "Sender identity and sales-desk recipient."],
    ["UPLOAD_DIR", "CMS attachments", "Private persistent storage path; defaults to ./data/uploads."],
    ["SHOPIFY_STORE_DOMAIN", "Required for Shopify", "Store domain used by Storefront and runtime Admin calls."],
    ["SHOPIFY_ADMIN_STORE_DOMAIN", "Collection tool only", "Permanent admin domain when different from storefront domain."],
    ["SHOPIFY_STOREFRONT_ACCESS_TOKEN", "Storefront", "Public Storefront token when configured."],
    ["SHOPIFY_STOREFRONT_PRIVATE_ACCESS_TOKEN", "Storefront", "Private Storefront token; server secret."],
    ["SHOPIFY_STOREFRONT_API_VERSION", "Optional", "Defaults to 2026-01 at runtime."],
    ["SHOPIFY_ADMIN_ACCESS_TOKEN", "Quotes/customer sync", "Required at runtime for draft orders, credit sync and consent changes."],
    ["SHOPIFY_ADMIN_API_VERSION", "Optional", "Runtime Admin API version; defaults to configured Storefront version or 2026-01."],
    ["SHOPIFY_STOREFRONT_PUBLICATION_ID", "Collection tool", "Needed when more than one possible Headless publication exists."],
    ["NODE_ENV / HOST / PORT", "Runtime", "Production is production / 0.0.0.0 / 80 in the container."],
]
table(doc, ["Variable", "Purpose", "Operational note"], config_rows, widths=[5, 4, 8], font_size=7.7)
callout(doc, "Secret handling", "Keep all production values runtime-only in Coolify or an approved secrets manager. Values prefixed VITE_ are browser-visible and must never contain secrets. Rotate any value suspected of exposure.", "danger")

heading(doc, "16. Backups, retention and recovery", 1)
para(doc, "The application persists two separate data sets: PostgreSQL records and the upload volume. Both must be backed up together. The repository does not define a business retention or deletion schedule; the client must approve one with legal/privacy advice.")
heading(doc, "16.1 Minimum backup procedure", 2)
bullets(doc, [
    "Schedule encrypted PostgreSQL backups with a documented retention period and off-host copy.",
    "Back up the cms_data upload volume on the same schedule and retain a mapping to the corresponding database backup.",
    "Protect backups with least-privilege access and test restoration at least quarterly.",
    "Record backup success/failure and assign an owner for remediation.",
    "Before a migration or high-risk deployment, take a verified pre-change backup and define rollback steps.",
])

heading(doc, "16.2 Recovery priorities", 2)
steps(doc, [
    "Stabilise the public storefront and prevent further writes if data consistency is at risk.",
    "Identify the required recovery point and restore PostgreSQL plus the matching upload snapshot.",
    "Start the app; allow idempotent migrations to run; check /ready.",
    "Verify CMS login, representative submissions, notes and attachment downloads.",
    "Verify Shopify connectivity separately; Shopify data is not restored from CMS backups.",
    "Document the incident, lost interval and follow-up controls.",
])

heading(doc, "17. Security and privacy operations", 1)
bullets(doc, [
    "Treat credit applications, bank details, trade references, customer contact information and attachments as confidential personal/business data.",
    "Use least-privilege access, strong unique passwords, MFA in Shopify/Coolify/email, and encrypted transport/storage/backups.",
    "Do not download attachments to unmanaged devices. Delete local copies according to the approved retention policy.",
    "The application uses CSRF protection for server functions, HTTP-only SameSite cookies, production Secure cookies, password hashing with scrypt, protected attachment reads, upload signature checks, rate limits and security headers.",
    "Public forms have a honeypot and an in-memory per-instance limit of eight requests per ten minutes per IP/scope. This is abuse reduction, not a complete WAF or distributed rate limit.",
    "Review legal/privacy wording, legal entity details, registered address, company/VAT numbers, retention periods and international-transfer disclosures with qualified UK counsel before launch.",
])

heading(doc, "18. Maintenance, testing and change control", 1)
table(doc, ["Command", "Purpose"], [
    ["npm run dev", "Local Vite development server."],
    ["npm run lint", "ESLint and formatting-rule verification."],
    ["npm run typecheck", "TypeScript verification without output."],
    ["npm run test", "Playwright end-to-end and unit-style browser tests."],
    ["npm run build", "Production application build."],
    ["npm run check (via Bun)", "Lint, typecheck, tests and build as a complete quality gate."],
    ["npx drizzle-kit generate", "Generate a database migration after an approved schema change."],
    ["npm run shopify:sync-collections", "Preview catalogue collection differences."],
], widths=[6, 11])

heading(doc, "18.1 Release acceptance", 2)
bullets(doc, [
    "No secrets in the commit or logs.",
    "Lint, typecheck, automated tests and production build pass.",
    "Database migration is reviewed, forward-safe and backed by a rollback/recovery plan.",
    "Desktop and mobile smoke tests cover navigation, search, catalogue, product, cart/quote and forms.",
    "CMS authentication and protected attachments remain inaccessible to unauthenticated users.",
    "Shopify publishing, checkout, quote draft-order creation and required email paths are tested in the target environment.",
    "Legal/content owners approve customer-facing wording changed by the release.",
])

heading(doc, "19. Troubleshooting", 1)
table(doc, ["Symptom", "Likely checks / response"], [
    ["Product missing", "Active? Headless publication? Correct manual collections? Variant available? Correct handle? Cache delay?"],
    ["Catalogue empty/error", "Confirm Storefront domain/token/API version and Shopify availability; inspect server logs. The site intentionally distinguishes integration failure from a genuine empty catalogue."],
    ["Quote fails", "Confirm Admin API domain/token/version and permissions for draft orders; inspect Shopify user errors and server logs."],
    ["Credit sync fails", "Confirm Admin token/scopes for customers and metafields; check duplicate/already-synced state and supplied email/data."],
    ["CMS unavailable", "Check DATABASE_URL and /ready; inspect PostgreSQL and migration logs."],
    ["Form saved but no email", "Check Resend/SMTP configuration, provider logs, MAIL_FROM verification and application logs. Process the CMS queue manually."],
    ["Attachment unavailable", "Check cms_data mount, UPLOAD_DIR, file ownership and whether database/files were restored together."],
    ["Cannot sign in to CMS", "Check APP_SESSION_SECRET, credentials and existing staff row. There is no self-service password reset."],
    ["Return customer did not receive update", "Check provider logs and address; status may still be saved. Contact manually and record a CMS note."],
    ["/health works but /ready fails", "Application process is alive; PostgreSQL is unavailable. Repair database connectivity without unnecessarily removing the storefront from the proxy."],
], widths=[4.5, 12.5], font_size=8.1)

heading(doc, "20. Launch and periodic checklists", 1)
heading(doc, "20.1 Before public launch", 2)
bullets(doc, [
    "Confirm production domain in site configuration, robots.txt and sitemap.xml.",
    "Obtain legal approval for terms, privacy, returns, delivery, cookies and disclaimer content; insert complete business identity details.",
    "Validate Shopify catalogue, Headless publication, Markets, shipping, tax, checkout, payments and notification emails with staging orders.",
    "Verify CMS seed account, session secret, database, persistent uploads, Resend/SMTP, sales-desk routing and backup schedules.",
    "Test all six main collections and every product-line collection with representative products.",
    "Test product resources, customer registration/login/reset, order history, quote creation, all CMS form types and protected attachments.",
    "Set named operational owners and response targets for New and In review submissions.",
])

heading(doc, "20.2 Weekly", 2)
bullets(doc, [
    "Clear or explain all New and In review CMS items.",
    "Review failed-email/application logs and Shopify integration errors.",
    "Spot-check key products, prices, availability, collection membership and resource links.",
    "Confirm backup jobs succeeded and storage capacity is healthy.",
])

heading(doc, "20.3 Monthly / quarterly", 2)
bullets(doc, [
    "Review Shopify and CMS access lists; remove stale access.",
    "Test restore procedure and attachment/database consistency.",
    "Review legal/contact details, category structure, broken links and customer-service copy.",
    "Review dependencies, API versions, security advisories and supported Node/PostgreSQL versions.",
    "Update this manual after any material operational change.",
])

page_break(doc)
heading(doc, "Appendix A — Route directory", 1)
route_rows = [
    ["/", "Homepage, category entry points, resources and part inquiry."],
    ["/products", "All products, category/product-line filters, availability, sort and load more."],
    ["/products/{handle}", "Product detail, variants, images, technical details, resources, cart and quote."],
    ["/search", "Global product search results."],
    ["/cart", "Shopify cart, quantities, totals, checkout and quote handoff."],
    ["/quote", "Browser quote list and Shopify draft-order submission."],
    ["/resources", "Product PDFs, manuals, datasheets and videos grouped by product type/category."],
    ["/login /register /forgot-password", "Shopify-backed customer access."],
    ["/account", "Customer details, credit-account guidance and Shopify order history."],
    ["/credit-account", "Credit account application stored in CMS."],
    ["/returns", "Return-request workflow; supports prefilled order reference."],
    ["/track-order", "Order-tracking request stored in CMS."],
    ["/contact-us /got-a-question", "Contact/product-question submission and attachment workflow."],
    ["/unsubscribe", "CMS unsubscribe request plus best-effort Shopify marketing-consent update."],
    ["/about-us", "Company information."],
    ["/terms-and-conditions", "Terms content."],
    ["/returns-policy", "Returns policy content."],
    ["/delivery-information", "Delivery information."],
    ["/privacy-policy", "Privacy information."],
    ["/cookies", "Cookie information and consent explanation."],
    ["/disclaimer", "Disclaimer content."],
    ["/admin/login", "Restricted CMS staff sign-in."],
    ["/admin", "CMS submission dashboard."],
    ["/admin/submissions/{id}", "Protected CMS detail, notes, attachments, status and credit sync."],
    ["/health", "Container liveness."],
    ["/ready", "PostgreSQL readiness."],
    ["/asphalt /concrete /packing /automation /home-controls /control-panels-software", "Legacy/convenience routes redirect to filtered /products views."],
    ["/new-arrivals", "Redirects to Control Panels & Software."],
]
table(doc, ["Route", "Purpose"], route_rows, widths=[6, 11], font_size=8)

heading(doc, "Appendix B — CMS data and workflow facts", 1)
table(doc, ["Fact", "Current implementation"], [
    ["Submission reference", "Human-friendly SA-{year}-{six-digit ID}."],
    ["Statuses", "new, in_review, approved, rejected, completed."],
    ["Dashboard page size", "25 records, newest first."],
    ["Admin session", "Encrypted/signed HTTP-only cookie, SameSite Lax, Secure in production, max age 8 hours."],
    ["Customer session", "Encrypted/signed HTTP-only cookie, SameSite Lax, Secure in production, max age 14 days subject to Shopify expiry."],
    ["Public rate limit", "8 attempts per 10 minutes per IP/scope in each running application instance."],
    ["Email durability", "Submission/status writes are not rolled back when provider delivery fails."],
    ["Attachments", "Database metadata plus private files in UPLOAD_DIR; authenticated retrieval only."],
    ["Database migrations", "Automatic, idempotent Drizzle migrations on boot; five retry attempts."],
], widths=[5, 12], font_size=8.2)

heading(doc, "Appendix C — Escalation information", 1)
para(doc, "The client should complete this table in its controlled copy. Do not place passwords or API keys here.")
table(doc, ["Area", "Named owner", "Contact / escalation window"], [
    ["Business owner", "________________", "________________"],
    ["Shopify catalogue", "________________", "________________"],
    ["Orders / sales desk", "________________", "________________"],
    ["Returns / CMS", "________________", "________________"],
    ["Coolify / deployment", "________________", "________________"],
    ["Database / backups", "________________", "________________"],
    ["Email provider", "________________", "________________"],
    ["Development support", "________________", "________________"],
], widths=[5, 5, 7])

heading(doc, "Appendix D — Known decisions still owned by the client", 1)
bullets(doc, [
    "Final legal entity, company number, VAT number, registered address and counsel-approved policies.",
    "Data retention/deletion schedule for submissions, bank/trade-reference details, notes, attachments, email and backups.",
    "Shopify shipping, tax, Markets, payments, fulfilment and customer-notification settings.",
    "Operational service levels, status definitions, escalation ownership and out-of-hours coverage.",
    "Whether the current CMS limitations require future staff management, export, deletion, audit history or resend features.",
])

callout(doc, "End of manual", "This document describes the implemented system as reviewed on 1 September 2026. Revalidate it after material code, Shopify, hosting or business-process changes.", "info")

# Prevent rows from splitting where possible and repeat header rows.
for tbl in doc.tables:
    if tbl.rows:
        tr_pr = tbl.rows[0]._tr.get_or_add_trPr()
        repeat = OxmlElement("w:tblHeader")
        repeat.set(qn("w:val"), "true")
        tr_pr.append(repeat)
    for row in tbl.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
